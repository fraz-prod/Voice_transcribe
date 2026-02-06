import docx
import sys
import os

def read_docx(path):
    if not os.path.exists(path):
        return f"Error: File not found {path}"
    doc = docx.Document(path)
    text = []
    
    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())
            
    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text.append(" | ".join(row_text))
            
    return "\n".join(text)

def main():
    filled_path = r"Day1_Visit_Filled_20260206 (1).docx"
    empty_path = r"[Internal] of Astria STAR 0215-301 Day 1.docx"
    
    print("--- FILLED DOC CONTENT ---")
    print(read_docx(filled_path))
    print("\n" + "="*50 + "\n")
    print("--- EMPTY DOC CONTENT ---")
    print(read_docx(empty_path))

if __name__ == "__main__":
    main()
