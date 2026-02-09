"""Check what's filled in the generated document"""
from docx import Document

doc = Document("test_gemini_filled.docx")

print("Searching for Notes section...")
notes_found = False
for i, para in enumerate(doc.paragraphs):
    if "Notes:" in para.text:
        print(f"Line {i}: {para.text[:100]}...")
        notes_found = True

if not notes_found:
    print("  -> Notes: label NOT found in document")

print("\nSearching for Injection sections...")
for i, para in enumerate(doc.paragraphs):
    if "Injection 1" in para.text or "Injection 2" in para.text:
        print(f"Line {i}: {para.text}")
        # Show next 10 lines
        for j in range(i+1, min(i+15, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                print(f"  Line {j}: {doc.paragraphs[j].text}")

print("\n" + "="*80)
print("Checking specific fields...")

keywords = ["Dose administered", "Laterality", "Start Date", "Start Time", "Notes"]
for keyword in keywords:
    found_count = 0
    for para in doc.paragraphs:
        if keyword in para.text:
            found_count += 1
            # Check if it has underscores (unfilled) or actual data
            if "___" in para.text or para.text.strip().endswith(":"):
                print(f"[UNFILLED] {keyword}: {para.text[:80]}...")
            else:
                print(f"[FILLED?] {keyword}: {para.text[:80]}...")
    
    if found_count == 0:
        print(f"[NOT FOUND] {keyword}")
