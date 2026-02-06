import docx
import sys
import os

def analyze_document(path):
    if not os.path.exists(path):
        print(f"Error: File not found {path}")
        return
    
    doc = docx.Document(path)
    
    print(f"\n{'='*60}")
    print(f"ANALYZING: {path}")
    print(f"{'='*60}")
    
    # Count paragraphs and tables
    print(f"\nTotal Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # Show table structures
    for i, table in enumerate(doc.tables):
        print(f"\n--- TABLE {i} ---")
        print(f"Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        for j, row in enumerate(table.rows):
            row_text = [cell.text.strip()[:30] for cell in row.cells]
            print(f"  Row {j}: {row_text}")
    
    # Show paragraphs with key markers
    print(f"\n--- KEY PARAGRAPHS ---")
    keywords = ["Weight", "Blood Pressure", "Heart Rate", "Temperature", "Respiratory", 
                "Time collected", "Date collected", "Start Time", "Start Date", 
                "Dose administered", "Laterality", "childbearing", "ECG", "PR (msec)"]
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(kw.lower() in text.lower() for kw in keywords):
            print(f"  [{i}] {text[:80]}...")

def main():
    # Analyze both empty and filled
    analyze_document("[Internal] of Astria STAR 0215-301 Day 1.docx")
    analyze_document("Day1_Visit_Filled_20260206 (2).docx")

if __name__ == "__main__":
    main()
