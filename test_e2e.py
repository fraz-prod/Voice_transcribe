"""
End-to-end test: Read sample.txt, extract data, fill form, analyze result
"""
import sys
import os
sys.path.append(os.getcwd())

from ai_services import LocalAIService
from form_filler import FormFiller
import docx

def run_e2e_test():
    print("="*60)
    print("END-TO-END FORM FILLING TEST")
    print("="*60)
    
    # 1. Read transcript
    print("\n[1] Reading sample.txt...")
    with open("sample.txt", "r") as f:
        transcript = f.read()
    print(f"    Transcript length: {len(transcript)} chars")
    
    # 2. Extract data
    print("\n[2] Extracting data...")
    data = LocalAIService.extract_data(transcript)
    
    print("\n    Key extracted fields:")
    print(f"    - visit_date: {data.get('visit_date')}")
    print(f"    - vitals_pre: {data.get('vitals_pre')}")
    print(f"    - vitals_post: {data.get('vitals_post')}")
    print(f"    - ecg: {data.get('ecg')}")
    print(f"    - labs: {data.get('labs')}")
    print(f"    - pregnancy: {data.get('pregnancy')}")
    print(f"    - injection: {data.get('injection')}")
    print(f"    - injection_2: {data.get('injection_2')}")
    
    # 3. Fill form
    print("\n[3] Filling form...")
    template = "[Internal] of Astria STAR 0215-301 Day 1.docx"
    filler = FormFiller(template)
    filled_doc_buffer = filler.fill_form(data, is_eligible=True)
    
    # Save to disk for inspection
    output_path = "test_filled_output.docx"
    with open(output_path, "wb") as f:
        f.write(filled_doc_buffer.read())
    print(f"    Saved to: {output_path}")
    
    # 4. Analyze filled form
    print("\n[4] Analyzing filled form...")
    doc = docx.Document(output_path)
    
    # Check tables
    print("\n    TABLE 0 (Pre-dose vitals):")
    if len(doc.tables) > 0:
        for row in doc.tables[0].rows:
            cell_text = row.cells[0].text.strip()
            print(f"      {cell_text}")
    
    print("\n    TABLE 1 (Post-dose vitals):")
    if len(doc.tables) > 1:
        for row in doc.tables[1].rows:
            cell_text = row.cells[0].text.strip()
            print(f"      {cell_text}")
    
    # Check key paragraphs
    print("\n    KEY PARAGRAPHS:")
    keywords = {
        "childbearing potential": None,
        "PR (msec)": None,
        "QRS (msec)": None,
        "Dose administered": [],
        "Start Time:": [],
        "Laterality:": [],
    }
    
    for para in doc.paragraphs:
        text = para.text.strip()
        for kw in keywords:
            if kw.lower() in text.lower():
                if isinstance(keywords[kw], list):
                    keywords[kw].append(text[:80])
                else:
                    keywords[kw] = text[:80]
    
    for kw, val in keywords.items():
        print(f"      {kw}: {val}")
    
    print("\n" + "="*60)
    print("TEST COMPLETE - Please open test_filled_output.docx to verify")
    print("="*60)

if __name__ == "__main__":
    run_e2e_test()
