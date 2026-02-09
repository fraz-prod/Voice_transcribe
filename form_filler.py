from docx import Document
from datetime import datetime
import io

class FormFiller:
    def __init__(self, template_path):
        self.template_path = template_path

    def fill_form(self, data: dict, is_eligible: bool = False) -> io.BytesIO:
        """
        Fills the DOCX template with data.
        Returns a BytesIO object of the filled document.
        """
        doc = Document(self.template_path)

        # Helper to mark Yes/No checkboxes
        def mark_yes_no(paragraph, answer_yes: bool):
            """
            Replaces ' Yes ' and ' No' patterns with [X] markers.
            answer_yes=True marks Yes, answer_yes=False marks No.
            """
            text = paragraph.text
            if answer_yes:
                text = text.replace("  Yes  ", "  [X] Yes  ").replace("  No", "  [ ] No")
                text = text.replace(" Yes  No", " [X] Yes  [ ] No")
            else:
                text = text.replace("  Yes  ", "  [ ] Yes  ").replace("  No", "  [X] No")
                text = text.replace(" Yes  No", " [ ] Yes  [X] No")
            paragraph.text = text

        # Helper to fill underscores
        def fill_underscores(paragraph, value: str):
            if not value:
                return
            text = paragraph.text
            # Replace various underscore patterns
            import re
            text = re.sub(r'_{3,}', str(value), text, count=1)
            paragraph.text = text
        
        # Helper to fill table cell
        def fill_table_cell(cell, value: str):
            if not value:
                return
            # Modify the first paragraph in the cell
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.text = para.text + " " + str(value)

        # === SECTION 1: Visit Performed ===
        for para in doc.paragraphs:
            if "Was the visit performed?" in para.text:
                mark_yes_no(para, True)
                break

        # === SECTION 2: Visit Date ===
        if "visit_date" in data:
            for para in doc.paragraphs:
                if "Visit Date:" in para.text:
                    fill_underscores(para, data["visit_date"])
                    break

        # === SECTION 3: Weekly Diary ===
        for para in doc.paragraphs:
            if "did the participant complete all weekly diary entries" in para.text:
                mark_yes_no(para, True)
                break

        # === SECTION 4: HAE Attack Contact ===
        for para in doc.paragraphs:
            if "did the site make contact with the participant to collect" in para.text:
                mark_yes_no(para, True)
                break

        # === SECTION 5: Adverse Events ===
        for para in doc.paragraphs:
            if "Were AEs reviewed at this visit?" in para.text:
                mark_yes_no(para, True)
                break

        for para in doc.paragraphs:
            if "Any new AEs reported?" in para.text:
                # Default: No new AEs (happy path)
                ae_status = data.get("adverse_events", "None")
                mark_yes_no(para, ae_status.lower() != "none")
                break

        # === SECTION 6: Concomitant Medications ===
        for para in doc.paragraphs:
            if "Were conmeds reviewed at this visit?" in para.text:
                mark_yes_no(para, True)
                break

        for para in doc.paragraphs:
            if "Any new or changes to conmeds reported?" in para.text:
                # Check if there are non-standard meds or changes
                meds = data.get("medications", [])
                has_changes = len(meds) > 1 or any(m.lower() != "takhzyro" for m in meds)
                mark_yes_no(para, has_changes)
                break

        # === SECTION 7: Eligibility Criteria ===
        for para in doc.paragraphs:
            if "Were all eligibility criteria met" in para.text:
                mark_yes_no(para, is_eligible)
                break

        for para in doc.paragraphs:
            if "Did the participant continue to meet eligibility at Day 1 visit?" in para.text:
                continued = data.get("continued_eligibility", "Yes")
                mark_yes_no(para, continued.lower() == "yes")
                break

        # === SECTION 8: IRT Randomization ===
        for para in doc.paragraphs:
            if "Subject randomized in IRT?" in para.text:
                mark_yes_no(para, True)
                break

        # === SECTION 9: Physical Exam (Pre-dose) ===
        for i, para in enumerate(doc.paragraphs):
            if "Physical Exam completed?" in para.text:
                mark_yes_no(para, True)
                break

        for para in doc.paragraphs:
            if "Any clinically significant abnormalities?" in para.text:
                mark_yes_no(para, False)  # Happy path: No abnormalities
                break

        # === SECTION 10: Pre-Dose Vitals (Table 0) ===
        if "vitals_pre" in data:
            vitals = data["vitals_pre"]
            
            # Try to find "Time collected" specifically for Pre-dose (before Labs)
            # Use a flag to ensure we only fill the first one we find in this context
            found_predose = False
            for para in doc.paragraphs:
                if "Physical Exam completed?" in para.text:
                    found_predose = True
                
                if found_predose and "Time collected" in para.text:
                    fill_underscores(para, vitals.get("time_collected", ""))
                    break
                
                if "Laboratory" in para.text: # Safety stop
                    break

            if len(doc.tables) > 0:
                table = doc.tables[0]
                try:
                    fill_table_cell(table.rows[0].cells[0], vitals.get('weight', ''))
                    fill_table_cell(table.rows[1].cells[0], vitals.get('bp', ''))
                    fill_table_cell(table.rows[2].cells[0], vitals.get('hr', ''))
                    fill_table_cell(table.rows[3].cells[0], vitals.get('temp', ''))
                    fill_table_cell(table.rows[4].cells[0], vitals.get('rr', ''))
                except IndexError:
                    pass

        # === SECTION 11: ECG ===
        for para in doc.paragraphs:
            if "Was the 12-lead ECG performed?" in para.text:
                mark_yes_no(para, True)
                break
        
        # Check for ECG Date specifically in ECG section
        ecg = data.get("ecg", {})
        found_ecg = False
        for para in doc.paragraphs:
            if "Was the 12-lead ECG performed?" in para.text:
                found_ecg = True
            
            if found_ecg:
                if "Date performed:" in para.text or "Day performed:" in para.text:
                    fill_underscores(para, ecg.get("date", ""))
                    break
                if "Laboratory" in para.text: 
                     break

        for para in doc.paragraphs:
            if "Results:" in para.text and "Normal" in para.text:
                result = ecg.get("result", "Normal")
                if result == "Normal":
                    para.text = para.text.replace(" Normal ", " [X] Normal ")
                break

        # === SECTION 12: Labs ===
        labs = data.get("labs", {})
        found_labs = False
        for para in doc.paragraphs:
            if "Laboratory assessments" in para.text or "Were all required labs collected?" in para.text:
                found_labs = True
                if "Were all required labs collected?" in para.text:
                     mark_yes_no(para, labs.get("collected", True))
            
            if found_labs:
                if "Date collected?" in para.text:
                    fill_underscores(para, labs.get("date", ""))
                elif "Time collected?" in para.text:
                    fill_underscores(para, labs.get("time", ""))
                elif "Urine collection time:" in para.text:
                    fill_underscores(para, labs.get("urine_time", ""))
                    break # Assuming this is the last field in Labs section we care about specific dates/times for

        for para in doc.paragraphs:
            if "Was Pharmacogenetics (DNA paxgene) sample collected?" in para.text:
                mark_yes_no(para, False)  # Often not collected
                break

        # === SECTION 13: Childbearing Potential ===
        pregnancy = data.get("pregnancy", {})
        for para in doc.paragraphs:
            if "Is subject of childbearing potential?" in para.text:
                mark_yes_no(para, pregnancy.get("potential", False))
                break
        
        # If childbearing potential is Yes, fill pregnancy test details
        if pregnancy.get("potential", False):
            for para in doc.paragraphs:
                if "was the sample collected?" in para.text.lower():
                    mark_yes_no(para, True)
                    break
            
            for para in doc.paragraphs:
                if "Collection Date:" in para.text:
                    fill_underscores(para, pregnancy.get("date", ""))
                    break
            
            for para in doc.paragraphs:
                if "Collection Time:" in para.text:
                    fill_underscores(para, pregnancy.get("time", ""))
                    break
            
            # Mark pregnancy result
            for para in doc.paragraphs:
                if "Results:" in para.text and ("Positive" in para.text or "Negative" in para.text):
                    result = pregnancy.get("result", "Negative")
                    if result == "Negative":
                        para.text = para.text.replace(" Negative", " [X] Negative").replace(" Positive", " [ ] Positive")
                    else:
                        para.text = para.text.replace(" Positive", " [X] Positive").replace(" Negative", " [ ] Negative")
                    break

        # === SECTION 14: Investigational Product ===
        for para in doc.paragraphs:
            if "Was investigational product administered?" in para.text:
                mark_yes_no(para, True)
                break

        # === SECTION 15: Injection 1 Details ===
        if "injection" in data:
            inj = data["injection"]
            for para in doc.paragraphs:
                if "Dose administered:" in para.text:
                    fill_underscores(para, inj.get("dose", ""))
                    break
            
            # Laterality checkbox marking
            for para in doc.paragraphs:
                if "Laterality:" in para.text:
                    lat = inj.get("laterality", "").lower()
                    if "left lower" in lat:
                        para.text = para.text.replace("Left Lower Quadrant", "[X] Left Lower Quadrant")
                    elif "left upper" in lat:
                        para.text = para.text.replace("Left Upper Quadrant", "[X] Left Upper Quadrant")
                    elif "right lower" in lat:
                        para.text = para.text.replace("Right Lower Quadrant", "[X] Right Lower Quadrant")
                    elif "right upper" in lat:
                        para.text = para.text.replace("Right Upper Quadrant", "[X] Right Upper Quadrant")
                    break
            
            for para in doc.paragraphs:
                if "Start Date:" in para.text:
                    fill_underscores(para, inj.get("start_date", ""))
                    break
            
            for para in doc.paragraphs:
                if "Start Time:" in para.text:
                    fill_underscores(para, inj.get("start_time", ""))
                    break

        # Injection 1 Interrupted
        for para in doc.paragraphs:
            if "Was injection interrupted?" in para.text:
                mark_yes_no(para, False)
                break

        # === SECTION 15b: Injection 2 Details ===
        if "injection_2" in data:
            inj2 = data["injection_2"]
            # Find Injection 2 section
            found_inj2 = False
            for i, para in enumerate(doc.paragraphs):
                if "Injection 2" in para.text:
                    found_inj2 = True
                    continue
                
                if found_inj2:
                    if "Dose administered:" in para.text:
                        fill_underscores(para, inj2.get("dose", ""))
                    elif "Laterality:" in para.text:
                        lat = inj2.get("laterality", "").lower()
                        if "left lower" in lat:
                            para.text = para.text.replace("Left Lower Quadrant", "[X] Left Lower Quadrant")
                        elif "right lower" in lat:
                            para.text = para.text.replace("Right Lower Quadrant", "[X] Right Lower Quadrant")
                        elif "left upper" in lat:
                            para.text = para.text.replace("Left Upper Quadrant", "[X] Left Upper Quadrant")
                        elif "right upper" in lat:
                            para.text = para.text.replace("Right Upper Quadrant", "[X] Right Upper Quadrant")
                    elif "Start Date:" in para.text:
                        fill_underscores(para, inj2.get("start_date", ""))
                    elif "Start Time:" in para.text:
                        fill_underscores(para, inj2.get("start_time", ""))
                    elif "Unblinded" in para.text:
                        break  # End of Injection 2 section

        # === SECTION 16: Post-Dose Vitals (Table 1) ===
        if "vitals_post" in data:
            vitals = data["vitals_post"]
            if len(doc.tables) > 1:
                table = doc.tables[1]
                try:
                    fill_table_cell(table.rows[0].cells[0], vitals.get('weight', ''))
                    fill_table_cell(table.rows[1].cells[0], vitals.get('bp', ''))
                    fill_table_cell(table.rows[2].cells[0], vitals.get('hr', ''))
                    fill_table_cell(table.rows[3].cells[0], vitals.get('temp', ''))
                    fill_table_cell(table.rows[4].cells[0], vitals.get('rr', ''))
                except IndexError:
                    pass

        # === SECTION 17: Physical Exam (Post-dose) ===
        # Find the SECOND occurrence of "Physical Exam completed?"
        pe_count = 0
        for para in doc.paragraphs:
            if "Physical Exam completed?" in para.text:
                pe_count += 1
                if pe_count == 2:
                    mark_yes_no(para, True)
                    break

        # Find the SECOND occurrence of abnormalities
        abn_count = 0
        for para in doc.paragraphs:
            if "Any clinically significant abnormalities?" in para.text:
                abn_count += 1
                if abn_count == 2:
                    mark_yes_no(para, False)
                    break

        # === ECG Section (Fill numeric fields) ===
        ecg = data.get("ecg", {})
        for para in doc.paragraphs:
            if "Heart rate (BPM):" in para.text:
                fill_underscores(para, ecg.get("hr", ""))
                break
        
        for para in doc.paragraphs:
            if "PR (msec):" in para.text:
                fill_underscores(para, ecg.get("pr", ""))
                break
        
        for para in doc.paragraphs:
            if "RR (msec):" in para.text:
                fill_underscores(para, ecg.get("rr", ""))
                break
        
        for para in doc.paragraphs:
            if "QRS (msec):" in para.text:
                fill_underscores(para, ecg.get("qrs", ""))
                break
        
        for para in doc.paragraphs:
            if "QT (msec):" in para.text:
                fill_underscores(para, ecg.get("qt", ""))
                break
        
        for para in doc.paragraphs:
            if "Date performed:" in para.text:
                fill_underscores(para, ecg.get("date", ""))
                break

        # === Notes Section ===
        if "notes" in data and data["notes"]:
            # Fill ALL Notes sections (there may be multiple)
            for para in doc.paragraphs:
                if "Notes:" in para.text:
                    # Use fill_underscores to replace the underscore pattern
                    fill_underscores(para, data["notes"])

        # === NEW: Overflow Information Section ===
        if "overflow_information" in data and data["overflow_information"]:
            overflow = data["overflow_information"]
            
            # Add section header
            doc.add_paragraph()
            header = doc.add_paragraph("Additional Medical Information (Beyond Protocol Fields)")
            header.runs[0].bold = True
            
            # Patient Concerns
            if overflow.get("patient_concerns"):
                p = doc.add_paragraph("Patient Concerns:")
                p.runs[0].bold = True
                for concern in overflow["patient_concerns"]:
                    doc.add_paragraph(f"  • {concern}")
            
            # Medication Questions
            if overflow.get("medication_questions"):
                p = doc.add_paragraph("Medication Questions:")
                p.runs[0].bold = True
                for question in overflow["medication_questions"]:
                    doc.add_paragraph(f"  • {question}")
            
            # Unreported Symptoms
            if overflow.get("unreported_symptoms"):
                p = doc.add_paragraph("Unreported Symptoms (mentioned but not as AE):")
                p.runs[0].bold = True
                for symptom in overflow["unreported_symptoms"]:
                    doc.add_paragraph(f"  • {symptom}")
            
            # Safety Observations
            if overflow.get("safety_observations"):
                p = doc.add_paragraph("Safety Observations:")
                p.runs[0].bold = True
                for obs in overflow["safety_observations"]:
                    doc.add_paragraph(f"  • {obs}")
            
            # Other Clinical Notes
            if overflow.get("other_clinical_notes"):
                p = doc.add_paragraph("Other Clinical Notes:")
                p.runs[0].bold = True
                for note in overflow["other_clinical_notes"]:
                    doc.add_paragraph(f"  • {note}")
        
        # === NEW: Validation Information ===
        if "validation" in data:
            val = data["validation"]
            doc.add_paragraph()
            val_header = doc.add_paragraph("Data Validation Summary")
            val_header.runs[0].bold = True
            
            completeness = val.get("completeness_score", "N/A")
            protocol_ok = "Yes" if val.get("protocol_compliance", False) else "No"
            overflow_detected = "Yes" if val.get("overflow_detected", False) else "No"
            requires_review = "Yes" if val.get("requires_review", False) else "No"
            
            doc.add_paragraph(f"Completeness Score: {completeness}%")
            doc.add_paragraph(f"Protocol Compliance: {protocol_ok}")
            doc.add_paragraph(f"Overflow Information Detected: {overflow_detected}")
            doc.add_paragraph(f"Requires Human Review: {requires_review}")
            
            if val.get("flags"):
                p = doc.add_paragraph("Flags/Warnings:")
                p.runs[0].bold = True
                for flag in val["flags"]:
                    doc.add_paragraph(f"  • {flag}")

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
